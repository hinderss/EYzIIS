import os
import re
import string
import time
import functools
from copy import deepcopy

import win32com.client as win32
import pythoncom

from alalyzer import analyzer

from docx import Document

from .models import Word


def convert_doc_to_docx(doc_path):
    pythoncom.CoInitialize()
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False
    doc = word.Documents.Open(doc_path)
    docx_path = os.path.splitext(doc_path)[0] + ".docx"
    doc.SaveAs(docx_path, FileFormat=16)  # 16 соответствует формату .docx
    doc.Close()
    word.Quit()
    return docx_path


def extract_text_from_docx(file_path):
    if file_path.endswith('.doc'):
        docx_path = convert_doc_to_docx(file_path)
        text = extract_text(docx_path)
        os.remove(docx_path)
        return text
    elif file_path.endswith('.docx'):
        return extract_text(file_path)
    else:
        raise ValueError("Unsupported file format")


def extract_text_with_styles(file_path):
    if file_path.path.endswith('.doc'):
        docx_path = convert_doc_to_docx(file_path)
        text = extract_text_style(docx_path)
        os.remove(docx_path)
        return text
    elif file_path.path.endswith('.docx'):
        return extract_text_style(file_path)
    else:
        raise ValueError("Unsupported file format")


def timing_decorator(func):
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.perf_counter()
        result = func(*args, **kwargs)
        end_time = time.perf_counter()
        elapsed_time = end_time - start_time
        print(f"Функция {func.__name__} выполнена за {elapsed_time:.6f} секунд")
        return result
    return wrapper


@timing_decorator
def analyze_text(text, file):
    words = re.findall(r"\b\w+\b", text.lower())

    unique_words = set(words)
    word_lexeme = [(word, analyzer.get_lexeme(word)) for word in unique_words]

    words_to_create = []
    for word, lexeme in word_lexeme:
        suffixes = [
            {
                "inflection": inflection,
                "suffix": suffix,
                "tag": tag,
                "explained": analyzer.explain_tag(tag),
            }
            for inflection, suffix, tag in lexeme.possible_suffixes
        ]

        word = Word(
            word=word,
            lemma=lexeme.lemma,
            pos=lexeme.pos,
            tag=lexeme.tag,
            suffixes=suffixes,
            file=file,
        )
        words_to_create.append(word)

    Word.objects.bulk_create(words_to_create)
    print("Кол-во слов", len(words_to_create))

    return Word.objects.filter(file=file)


def extract_text(file_path):
    doc = Document(file_path)
    text = " ".join([p.text for p in doc.paragraphs])
    return text


def extract_text_style(file_path):
    doc = Document(file_path)
    html_output = ""

    for para in doc.paragraphs:
        alignment = "left"  # По умолчанию
        if para.alignment == 1:
            alignment = "center"
        elif para.alignment == 2:
            alignment = "right"

        html_output += f'<p style="text-align: {alignment};">'

        for run in para.runs:
            font_size = run.font.size
            if font_size:
                font_size_pt = font_size.pt

            font_color = run.font.color
            if font_color and font_color.rgb:
                color = font_color.rgb

            words = run.text.split()

            for word in words:
                html_output += f'<span>'

                if font_size:
                    html_output += f'<span style="font-size: {font_size_pt}pt;">'
                if font_color and font_color.rgb:
                    html_output += f'<span style="color: #{color[2:]}">'

                if run.bold:
                    html_output += "<b>"
                if run.italic:
                    html_output += "<i>"
                if run.underline:
                    html_output += "<u>"

                initial_word = deepcopy(word)
                cleaned_word = word.translate(str.maketrans('', '', string.punctuation))
                if cleaned_word:
                    lexeme = analyzer.get_lexeme(cleaned_word)
                    suffixes = [f"-{suffix}: {analyzer.explain_tag(tag)}" for _, suffix, tag in lexeme.possible_suffixes]
                    html_output += (f'<span data-bs-toggle="tooltip" title="'
                                    f'{lexeme.lemma}: '
                                    f'Основа: {lexeme.stem}, '
                                    f'Часть речи: {analyzer.explain_pos(lexeme.pos)}, '
                                    f'Окончания: {"; ".join(map(str, suffixes))}'
                                    f'">{cleaned_word} </span>')
                else:
                    html_output += f'<span>{initial_word}</span>'
                if run.underline:
                    html_output += "</u>"
                if run.italic:
                    html_output += "</i>"
                if run.bold:
                    html_output += "</b>"

                if font_color and font_color.rgb:
                    html_output += "</span>"

                if font_size:
                    html_output += "</span>"

                html_output += "</span>"

        html_output += "</p>"
    return html_output
